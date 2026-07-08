#!/usr/bin/env python3
"""職務経歴書Markdown→Word変換（原本テンプレート書式を適用・バッジ集約版）"""
import re
from docx import Document
from docx.oxml.shared import OxmlElement, qn

MD_PATH = "/home/yn4416/projects/obsidian-ssot/40_CAREER/01_ドキュメント/職務経歴書_改善版_2026-07-08.md"
TPL_PATH = "/tmp/template_original.docx"
OUT_PATH = "/tmp/resume_formatted.docx"


def add_hyperlink(paragraph, url, text):
    """段落にクリック可能ハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs_with_bold(paragraph, text):
    """**太字** を処理してrunに分ける。"""
    segments = re.split(r"(\*\*[^*]+\*\*)", text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
        else:
            paragraph.add_run(seg)


def extract_badge_detail(alt, img_url):
    """shields.io 画像URLから意味ある詳細テキストを抽出。
    例: /badge/python-3.12+-blue → 'Python 3.12+'"""
    if "shields.io" in img_url:
        m = re.search(r"/badge/(.+)$", img_url)
        if m:
            raw = m.group(1)
            parts = raw.split("-")
            if len(parts) >= 3:
                label = parts[0].capitalize()
                message = "-".join(parts[1:-1])
            elif len(parts) == 2:
                label = parts[0].capitalize()
                message = parts[1]
            else:
                return alt
            message = message.replace("%20", " ").replace("%2C", ",")
            return f"{label} {message}" if message else label
    return alt


def parse_badges(line):
    """1行からバッジ (detail, link_url) のリストを抽出。
    URL無し/相対URLにも対応。"""
    badges = []
    pattern = r"\[!\[([^\]]*)\]\(([^)]*)\)\](?:\(([^)]*)\))?"
    for m in re.finditer(pattern, line):
        alt = m.group(1).strip()
        img_url = m.group(2).strip()
        link_url = m.group(3).strip() if m.group(3) else ""
        if link_url and not link_url.startswith("http"):
            link_url = ""
        detail = extract_badge_detail(alt, img_url)
        badges.append((detail, link_url))
    return badges


def main():
    doc = Document(TPL_PATH)

    # 全段落・全表クリア（sectPrは残す）
    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if tag.endswith("}sectPr"):
            continue
        if tag.endswith("}p") or tag.endswith("}tbl"):
            body.remove(child)

    # Markdown読込・frontmatterスキップ
    with open(MD_PATH, encoding="utf-8") as f:
        raw_lines = f.read().split("\n")

    fm_delim_count = 0
    content_lines = []
    for line in raw_lines:
        if line.strip() == "---":
            fm_delim_count += 1
            if fm_delim_count <= 2:
                continue
        if fm_delim_count >= 2:
            content_lines.append(line)

    i = 0
    while i < len(content_lines):
        line = content_lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 区切り線 ---
        if stripped == "---":
            sep = doc.add_paragraph("────────────────────────────────────────")
            sep.style = doc.styles["Normal"]
            i += 1
            continue

        # 見出し（絵文字除去: ビジネス文書向け）
        if stripped.startswith("### "):
            title = stripped[4:].strip()
            title = re.sub(r"^[🔹▸▶▪♦✦▸▸▸\s]+", "", title).strip()
            doc.add_heading(title, level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # 連続するバッジ行([![ で始まる)を集約 → 1つのList Bullet
        if stripped.startswith("[!["):
            all_badges = []
            while i < len(content_lines) and content_lines[i].strip().startswith("[!["):
                all_badges.extend(parse_badges(content_lines[i].strip()))
                i += 1
            if all_badges:
                p = doc.add_paragraph(style="List Bullet")
                labels = [b[0] for b in all_badges]
                run = p.add_run(f"公開バッジ: {' / '.join(labels)}")
            continue

        # 🔗 URL 行 → GitHub リンク
        if stripped.startswith("🔗"):
            url = stripped.replace("🔗", "").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("GitHub: ")
            add_hyperlink(p, url, url)
            i += 1
            continue

        # 箇条書き - (太字ラベル付き)
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, text)
            i += 1
            continue

        # 表 | ... |
        if stripped.startswith("|"):
            table_rows = []
            while i < len(content_lines) and content_lines[i].strip().startswith("|"):
                row_line = content_lines[i].strip()
                if re.match(r"^\|[\s:|-]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(table_rows):
                    for ci in range(ncols):
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_runs_with_bold(para, row[ci] if ci < len(row) else "")
                doc.add_paragraph()
            continue

        # 通常段落
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(f"✅ 変換完了: {OUT_PATH}")

    doc2 = Document(OUT_PATH)
    print(f"\n=== 生成docx 構造（全段落）===")
    for p in doc2.paragraphs:
        if p.text.strip():
            print(f"  [{p.style.name:12s}] {p.text[:70]}")
    print(f"\n表数: {len(doc2.tables)}")
    for ti, tbl in enumerate(doc2.tables):
        print(f"  表{ti+1}: {len(tbl.rows)}行 x {len(tbl.columns)}列")
        for ri, row in enumerate(tbl.rows):
            cells = [c.text[:20] for c in row.cells]
            print(f"    行{ri}: {cells}")


if __name__ == "__main__":
    main()
